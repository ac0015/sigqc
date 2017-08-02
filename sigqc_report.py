from docx import Document
from docx.shared import Inches
import types

################################################################################################
# sigqc_report.py
# Austin Coleman
#
# Module to automatically create Word Documents for
# SigQC results using SigQCReport objects.
#
# Example Usage:
#   fig_to_save = "My Figure Path\\My Figure Name"
#   report = SigQCReport()
#
#   ### Have report save in current directory ###
#   report.setFilePath("")
#   report.setDocName("MyTestDocument")
#
#   ### Add a Section to the document with figures ###
#   report.addSection(i_title="My Header", i_description="Some Text", i_figures=fig_to_save)
#   report.writeReport()
#
################################################################################################

##########################
# SigQCReport Class
##########################
class SigQCReport:
    '''
    The SigQCReport class is the base class from which a variety of SigQC data reports may
    be created. 
    '''
    def __init__(self, template_name="SigQCPCATemplate.docx"):
        '''
        Constructor of the SigQCReport class used to export SigQC Product Data into a Word
        Document format. The optional variable template_name is used to specify
        which Document template to open and which to save as an external file.
        '''
        self._fpath = None
        self._document = Document(template_name)
        self._docname = None
        
        # Set narrow margins to fit larger figures
        sections = self._document.sections
        section = sections[0]
        self._section = section
        self._section.left_margin = Inches(0.8)
        self._section.right_margin = Inches(0.8)
    
    def setFilePath(self, i_filepath): 
        '''
        Takes a string specifying the location of the Word document to be created from
        the SigQCReport object.
        '''
        self._fpath = i_filepath
    
    def setDocName(self, i_filename):
        '''
        Takes a string specifying the name of the Word document to be saved by the SigQCReport
        object.
        '''
        if (".doc" in i_filename):
            self._docname = i_filename
        else:
            self._docname = i_filename + ".docx"
    
    def addSection(self, i_title, i_description = None, i_figures = None):
        '''
        Adds a section within the Word document of the SigQCReport object underneath the specified
        title.
        
        Inputs
        ------
            i_title - String describing the heading of the section of the Word document
            i_description - (Optional) Any string-like object that can be written to a
                document that will precede the figures.
            i_figures - (Optional) A string or list of strings specifying the filepath and filename
                of any figures to be added to the document.
        '''
        self._document.add_heading(i_title)
        if (i_description != None):
            self._document.add_paragraph(i_description)
        if (i_figures != None):
            if not isinstance(i_figures, str):
                for figure in i_figures:
                    self._document.add_picture(figure)
            else:
                self._document.add_picture(i_figures)
    
    def writeReport(self, o_filepath = None, o_docname = None):
        '''
        Saves the Word document of the SigQCReport
        
        Inputs
        ------
        o_filepath - (Optional) If not already specified, takes a string and will be used 
            as the location to store the Word document
        o_docname - (Optional) If not already specified, takes a string and will be used as
            the name of the Word document

        Outputs
        -------
            Saves the Word Document to the filepath of the object.
            Does not return anything explicitly.
        '''
        if (i_filepath != None):
            self.setFilePath(o_filepath)
        elif (self._fpath == None):
            self.setFilePath("")
        if (i_docname != None):
            self.setDocName(o_docname)
        elif (self._docname == None):
            self.setDocName("SigQCReportDoc.docx")
        self._document.save(self._fpath+self._docname)
